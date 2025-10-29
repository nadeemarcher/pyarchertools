from flask import Flask, render_template, request, send_file, jsonify
from werkzeug.utils import secure_filename
from docx import Document
from docx.enum.text import WD_BREAK
from docxcompose.composer import Composer
from datetime import datetime
import os
import io
import zipfile
import shutil
import traceback

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB
app.config['UPLOAD_FOLDER'] = 'temp_uploads'

# Ensure upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/text-formatter')
def text_formatter():
    return render_template('text_formatter.html')

@app.route('/docx-cover-merger')
def docx_cover_merger():
    return render_template('docx_cover_merger.html')

def allowed_file(filename: str) -> bool:
    return bool(filename) and filename.lower().endswith('.docx')

@app.route('/api/merge-docx', methods=['POST'])
def merge_docx():
    uploaded_paths = []
    merged_paths = []
    try:
        if 'template' not in request.files or 'content_files[]' not in request.files:
            return jsonify({'error': 'Missing files'}), 400

        template_file = request.files['template']
        content_files = request.files.getlist('content_files[]')

        if not template_file or not content_files:
            return jsonify({'error': 'No files provided'}), 400

        if not allowed_file(template_file.filename):
            return jsonify({'error': 'Template must be a .docx file'}), 400

        # Save template (unique name)
        ts = datetime.now().strftime('%Y%m%d%H%M%S%f')
        template_name = f"{ts}_{secure_filename(template_file.filename)}"
        template_path = os.path.join(app.config['UPLOAD_FOLDER'], template_name)
        template_file.save(template_path)
        uploaded_paths.append(template_path)

        # Merge each content file (fresh template usage inside helper)
        for cf in content_files:
            if not allowed_file(cf.filename):
                continue
            content_name = f"{datetime.now().strftime('%Y%m%d%H%M%S%f')}_{secure_filename(cf.filename)}"
            content_path = os.path.join(app.config['UPLOAD_FOLDER'], content_name)
            cf.save(content_path)
            uploaded_paths.append(content_path)

            out_basename = secure_filename(cf.filename).rsplit('.docx', 1)[0] + '_with_cover.docx'
            merged_path = merge_documents_properly(template_path, content_path, out_basename)
            merged_paths.append({'filename': out_basename, 'path': merged_path})

        if not merged_paths:
            cleanup_files(uploaded_paths)
            return jsonify({'error': 'No valid .docx content files uploaded'}), 400

        # Read merged results into memory, delete all files immediately, then return in-memory response
        if len(merged_paths) == 1:
            m = merged_paths[0]
            with open(m['path'], 'rb') as f:
                data = f.read()
            # delete on-disk files now
            cleanup_files(uploaded_paths + [mp['path'] for mp in merged_paths])
            bio = io.BytesIO(data)
            bio.seek(0)
            return send_file(
                bio,
                as_attachment=True,
                download_name=m['filename'],
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

        # multiple -> create in-memory zip
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
            for m in merged_paths:
                with open(m['path'], 'rb') as f:
                    file_bytes = f.read()
                zf.writestr(m['filename'], file_bytes)
        zip_buffer.seek(0)
        # delete disk files immediately
        cleanup_files(uploaded_paths + [mp['path'] for mp in merged_paths])
        return send_file(
            zip_buffer,
            as_attachment=True,
            download_name='merged_documents.zip',
            mimetype='application/zip'
        )

    except Exception as e:
        traceback.print_exc()
        # attempt cleanup on error
        try:
            cleanup_files(uploaded_paths + [mp['path'] for mp in merged_paths])
        except Exception:
            pass
        return jsonify({'error': str(e)}), 500

def merge_documents_properly(template_path: str, content_path: str, output_filename: str) -> str:
    """
    Use the template as the base document so its styles, headers and footers are preserved.
    Add a single page break after the template, then append the content.
    """
    # Load template as the base so headers/footers/styles are kept
    template_doc = Document(template_path)
    composer = Composer(template_doc)

    # Ensure single page break after the template (cover)
    if template_doc.paragraphs:
        last_par = template_doc.paragraphs[-1]
    else:
        last_par = template_doc.add_paragraph()
    # Add a page break run if not already the last element being a page break.
    # Adding a run with a page break ensures content starts on the next page.
    last_par.add_run().add_break(WD_BREAK.PAGE)

    # Append content document
    composer.append(Document(content_path))

    out_path = os.path.join(
        app.config['UPLOAD_FOLDER'],
        f"merged_{datetime.now().strftime('%Y%m%d%H%M%S%f')}_{secure_filename(output_filename)}"
    )
    composer.save(out_path)
    return out_path

def cleanup_files(paths):
    """Remove all given file paths and ignore errors."""
    for p in paths:
        try:
            if not p:
                continue
            if os.path.isdir(p):
                shutil.rmtree(p, ignore_errors=True)
            else:
                if os.path.exists(p):
                    os.remove(p)
        except Exception:
            pass

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)